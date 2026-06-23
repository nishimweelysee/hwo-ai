#!/usr/bin/env python3
"""Generate HWO project documentation Word document with improved diagrams and explanations."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "docs" / "project-doc"
DIAGRAM_DIR = DOC_DIR / "diagrams"
SCREENSHOT_DIR = DOC_DIR / "screenshots"
OUTPUT = DOC_DIR / "Health_Workforce_Optimizer_Project_Documentation.docx"

# ── Visual constants ──────────────────────────────────────────────────────────
C_PRIMARY = "#1565C0"
C_SECONDARY = "#37474F"
C_ACCENT = "#00897B"
C_ARROW = "#263238"
C_LABEL = "#455A64"
C_ENTITY = "#E3F2FD"
C_ENTITY_BORDER = "#1565C0"
C_LAYER_FE = "#E3F2FD"
C_LAYER_BE = "#E8F5E9"
C_LAYER_AI = "#FFF3E0"
C_LAYER_DB = "#F3E5F5"


def save_fig(name: str, dpi: int = 200):
    path = DIAGRAM_DIR / name
    plt.tight_layout()
    plt.savefig(path, dpi=dpi, bbox_inches="tight", facecolor="white", pad_inches=0.15)
    plt.close()
    return path


# ── Diagram helpers ───────────────────────────────────────────────────────────

def draw_arrow(ax, x1, y1, x2, y2, label="", color=C_ARROW, style="-|>", lw=1.8,
               label_offset=(0, 0.08), dashed=False, rad=0.0):
    """Draw a precise arrow with optional curved routing."""
    ls = (0, (5, 3)) if dashed else "solid"
    connection = f"arc3,rad={rad}" if rad else "arc3,rad=0"
    arrow = FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle=style,
        color=color,
        linewidth=lw,
        linestyle=ls,
        mutation_scale=14,
        connectionstyle=connection,
        shrinkA=2,
        shrinkB=2,
        zorder=3,
    )
    ax.add_patch(arrow)
    if label:
        mx, my = (x1 + x2) / 2 + label_offset[0], (y1 + y2) / 2 + label_offset[1]
        ax.text(mx, my, label, ha="center", va="bottom", fontsize=7.5,
                color=C_LABEL, fontweight="medium",
                bbox=dict(boxstyle="round,pad=0.15", facecolor="white", edgecolor="#CFD8DC", alpha=0.95))


def draw_orthogonal(ax, points, label="", color=C_ARROW, dashed=False):
    """Draw orthogonal polyline with arrowhead at end."""
    ls = (0, (4, 3)) if dashed else "solid"
    for i in range(len(points) - 2):
        ax.plot([points[i][0], points[i + 1][0]], [points[i][1], points[i + 1][1]],
                color=color, lw=1.8, solid_capstyle="round", zorder=2)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    arrow = FancyArrowPatch(
        (x1, y1), (x2, y2), arrowstyle="-|>", color=color, linewidth=1.8,
        linestyle=ls, mutation_scale=14, shrinkA=0, shrinkB=2, zorder=3,
    )
    ax.add_patch(arrow)
    if label:
        mid = points[len(points) // 2]
        ax.text(mid[0], mid[1] + 0.12, label, ha="center", fontsize=7.5, color=C_LABEL,
                bbox=dict(boxstyle="round,pad=0.12", facecolor="white", edgecolor="#CFD8DC", alpha=0.95))


def draw_crows_foot(ax, x, y, direction="right"):
    """Draw crow's foot notation for ERD 'many' side."""
    size = 0.12
    if direction == "right":
        ax.plot([x, x + size], [y, y + size * 0.8], color=C_ARROW, lw=1.2)
        ax.plot([x, x + size], [y, y], color=C_ARROW, lw=1.2)
        ax.plot([x, x + size], [y, y - size * 0.8], color=C_ARROW, lw=1.2)
    elif direction == "left":
        ax.plot([x, x - size], [y, y + size * 0.8], color=C_ARROW, lw=1.2)
        ax.plot([x, x - size], [y, y], color=C_ARROW, lw=1.2)
        ax.plot([x, x - size], [y, y - size * 0.8], color=C_ARROW, lw=1.2)
    elif direction == "down":
        ax.plot([x - size * 0.8, x, x + size * 0.8], [y, y - size, y - size], color=C_ARROW, lw=1.2)


def draw_entity(ax, x, y, name, fields, w=2.0):
    """Draw ERD entity box with header and attribute list."""
    header_h = 0.38
    row_h = 0.26
    body_h = len(fields) * row_h + 0.1
    h = header_h + body_h
    left, bottom = x - w / 2, y - h / 2

    # Body
    body = FancyBboxPatch((left, bottom), w, body_h, boxstyle="square,pad=0",
                          facecolor="white", edgecolor=C_ENTITY_BORDER, linewidth=1.5, zorder=1)
    ax.add_patch(body)
    # Header
    header = FancyBboxPatch((left, bottom + body_h), w, header_h, boxstyle="square,pad=0",
                            facecolor=C_ENTITY_BORDER, edgecolor=C_ENTITY_BORDER, linewidth=1.5, zorder=2)
    ax.add_patch(header)
    ax.text(x, bottom + body_h + header_h / 2, name, ha="center", va="center",
            fontsize=9, fontweight="bold", color="white", zorder=3)

    for i, field in enumerate(fields):
        is_pk = "PK" in field
        is_fk = "FK" in field
        weight = "bold" if is_pk else "normal"
        color = C_PRIMARY if is_pk else ("#6A1B9A" if is_fk else "#333")
        prefix = "* " if is_pk else ("-> " if is_fk else "  ")
        ax.text(left + 0.08, bottom + body_h - 0.18 - i * row_h, prefix + field,
                ha="left", va="top", fontsize=7, fontweight=weight, color=color, zorder=3)

    return {"x": x, "y": y, "w": w, "h": h, "left": left, "right": left + w,
            "top": bottom + h, "bottom": bottom, "cx": x, "cy": y}


def edge_point(box, target_box):
    """Return edge point on box facing target_box."""
    dx = target_box["cx"] - box["cx"]
    dy = target_box["cy"] - box["cy"]
    if abs(dx) >= abs(dy):
        return (box["right"] if dx > 0 else box["left"], box["cy"])
    return (box["cx"], box["top"] if dy > 0 else box["bottom"])


def draw_actor(ax, x, y, name):
    """UML stick-figure actor."""
    ax.plot(x, y + 0.35, "o", color=C_SECONDARY, markersize=10, zorder=4)
    ax.plot([x, x], [y + 0.1, y - 0.15], color=C_SECONDARY, lw=2, zorder=4)
    ax.plot([x - 0.22, x + 0.22], [y - 0.45, y - 0.45], color=C_SECONDARY, lw=2, zorder=4)
    ax.text(x, y - 0.72, name, ha="center", va="top", fontsize=8, fontweight="medium", color=C_SECONDARY)


def draw_use_case(ax, x, y, label, w=2.1, h=0.55):
    e = mpatches.Ellipse((x, y), w, h, facecolor=C_LAYER_FE, edgecolor=C_PRIMARY, linewidth=1.5, zorder=2)
    ax.add_patch(e)
    ax.text(x, y, label, ha="center", va="center", fontsize=7.5, color=C_SECONDARY, zorder=3)
    return {"x": x, "y": y, "w": w, "h": h}


# ── Diagrams ──────────────────────────────────────────────────────────────────

def draw_erd():
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis("off")
    ax.set_title("Entity Relationship Diagram (ERD) — PostgreSQL Schema", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    # Row 1 — Core identity
    dept = draw_entity(ax, 2.5, 9.2, "Department", [
        "id  PK", "name", "code", "staffCount", "workload", "organizationId",
    ])
    staff = draw_entity(ax, 6.0, 9.2, "Staff", [
        "id  PK", "name", "email", "role", "departmentId  FK",
    ])
    user = draw_entity(ax, 9.5, 9.2, "User", [
        "id  PK", "email", "password", "role", "staffId", "organization",
    ])
    profile = draw_entity(ax, 13.0, 9.2, "UserProfile", [
        "id  PK", "userId  FK", "preferences (JSON)", "dashboardSettings",
    ])

    # Row 2 — Operations
    schedule = draw_entity(ax, 2.5, 6.2, "Schedule", [
        "id  PK", "staffId  FK", "departmentId  FK", "date", "shift", "status",
    ])
    workload = draw_entity(ax, 6.0, 6.2, "WorkloadRecord", [
        "id  PK", "departmentId  FK", "date", "hour", "patientVolume", "workload",
    ])
    wellness = draw_entity(ax, 9.5, 6.2, "WellnessRecord", [
        "id  PK", "staffId  FK", "score", "overtimeHours", "recordedAt",
    ])
    cert = draw_entity(ax, 13.0, 6.2, "Certification", [
        "id  PK", "staffId  FK", "name", "issuer", "expiryDate", "status",
    ])

    # Row 3 — Support
    resource = draw_entity(ax, 2.5, 3.2, "Resource", [
        "id  PK", "name", "sku", "quantity", "reorderLevel", "location",
    ])
    procure = draw_entity(ax, 6.0, 3.2, "ProcurementRequest", [
        "id  PK", "resourceId  FK", "quantity", "status", "requestedBy",
    ])
    audit = draw_entity(ax, 9.5, 3.2, "AuditLog", [
        "id  PK", "userId  FK", "action", "module", "timestamp", "details",
    ])
    model = draw_entity(ax, 13.0, 3.2, "PredictionModel", [
        "id  PK", "name", "algorithm", "accuracy", "trainedAt", "artifactPath",
    ])

    relations = [
        (dept, staff, "1", "N"),
        (dept, workload, "1", "N"),
        (dept, schedule, "1", "N"),
        (staff, schedule, "1", "N"),
        (staff, wellness, "1", "N"),
        (staff, cert, "1", "N"),
        (user, profile, "1", "1"),
        (user, audit, "1", "N"),
        (resource, procure, "1", "N"),
    ]

    for a, b, card_a, card_b in relations:
        p1 = edge_point(a, b)
        p2 = edge_point(b, a)
        draw_arrow(ax, p1[0], p1[1], p2[0], p2[1], label=f"{card_a}:{card_b}", lw=1.5)
        # Crow's foot on many side
        if card_b == "N":
            draw_crows_foot(ax, p2[0], p2[1], "left" if p2[0] < b["cx"] else "right")
        if card_a == "N":
            draw_crows_foot(ax, p1[0], p1[1], "right" if p1[0] > a["cx"] else "left")

    ax.text(8, 0.5,
            "Notation: PK = Primary Key  |  FK = Foreign Key  |  1:N = One-to-Many  |  1:1 = One-to-One",
            ha="center", fontsize=8, color=C_LABEL, style="italic")
    return save_fig("erd.png")


def draw_class_diagram():
    fig, ax = plt.subplots(figsize=(15, 10))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.set_title("Class Diagram — Spring Boot Layered Architecture", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    def draw_layer(y, title, classes, color, border):
        frame = FancyBboxPatch((0.4, y - 0.55), 14.2, 1.35, boxstyle="round,pad=0.06",
                               facecolor=color, edgecolor=border, linewidth=2, alpha=0.35, zorder=0)
        ax.add_patch(frame)
        ax.text(0.7, y + 0.55, title, fontsize=10, fontweight="bold", color=border)
        x = 0.8
        boxes = []
        for cls in classes:
            w = max(1.55, len(cls) * 0.085 + 0.4)
            box = FancyBboxPatch((x, y - 0.25), w, 0.55, boxstyle="round,pad=0.03",
                                 facecolor="white", edgecolor=border, linewidth=1.3, zorder=2)
            ax.add_patch(box)
            ax.text(x + w / 2, y + 0.025, cls, ha="center", va="center", fontsize=7.2, zorder=3)
            boxes.append((x + w / 2, y + 0.025))
            x += w + 0.12
        return boxes

    c1 = draw_layer(8.2, "«presentation»  Controller Layer", [
        "AuthController", "ScheduleController", "PredictionsController",
        "WellnessController", "ResourceController", "ComplianceController", "ImportController",
    ], C_LAYER_FE, C_PRIMARY)

    c2 = draw_layer(5.8, "«business»  Service Layer", [
        "UserService", "SchedulingService", "PredictionService", "WellnessService",
        "ResourceService", "ImportService", "ComplianceService", "AiServiceClient",
    ], C_LAYER_BE, "#2E7D32")

    c3 = draw_layer(3.4, "«persistence»  Repository Layer (JPA)", [
        "StaffRepository", "ScheduleRepository", "WorkloadRecordRepository",
        "UserRepository", "ResourceRepository", "AuditLogRepository", "ComplianceRecordRepository",
    ], C_LAYER_DB, "#6A1B9A")

    c4 = draw_layer(1.0, "«domain»  Entity Layer (@Entity)", [
        "Staff", "Department", "Schedule", "User", "WorkloadRecord",
        "Resource", "WellnessRecord", "Certification", "ComplianceRecord",
    ], "#FFFDE7", "#F9A825")

    # Vertical dependency arrows between layers
    for y_from, y_to in [(7.55, 6.45), (5.15, 4.05), (2.75, 1.65)]:
        draw_arrow(ax, 7.5, y_from, 7.5, y_to, color=C_ACCENT, lw=2.2, style="-|>")
        draw_arrow(ax, 7.5, y_to + 0.05, 7.5, y_from - 0.05, color=C_LABEL, lw=1.2,
                   style="-|>", dashed=True, label_offset=(0.35, 0))

    ax.text(7.5, 0.35,
            "Dependency flow: Controller → Service → Repository → Entity  (solid = calls, dashed = returns DTO/entity)",
            ha="center", fontsize=8, color=C_LABEL, style="italic")
    return save_fig("class_diagram.png")


def draw_usecase():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Use Case Diagram — Health Workforce Optimizer", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    boundary = FancyBboxPatch((3.2, 0.6), 7.6, 7.5, boxstyle="round,pad=0.12",
                              facecolor="#FAFAFA", edgecolor=C_SECONDARY, linewidth=2.5, zorder=0)
    ax.add_patch(boundary)
    ax.text(7.0, 7.85, "«system»  Health Workforce Optimizer (HWO)", ha="center",
            fontsize=11, fontweight="bold", color=C_SECONDARY)

    ucs = {
        "View Dashboard": (4.5, 6.8),
        "Import Workforce Data": (7.0, 6.8),
        "Analyze Workload": (9.5, 6.8),
        "Train AI Model": (4.5, 5.5),
        "Manage Schedules": (7.0, 5.5),
        "Generate Reports": (9.5, 5.5),
        "Monitor Wellness": (4.5, 4.2),
        "Manage Inventory": (7.0, 4.2),
        "Track Compliance": (9.5, 4.2),
        "Configure System": (5.75, 2.9),
        "Audit Activity": (8.25, 2.9),
        "Mobile Schedule View": (7.0, 1.5),
    }
    uc_boxes = {k: draw_use_case(ax, x, y, k) for k, (x, y) in ucs.items()}

    actors = {
        "Admin": (1.2, 6.5),
        "Manager": (1.2, 5.0),
        "Analyst": (1.2, 3.5),
        "Scheduler": (1.2, 2.0),
        "Healthcare\nStaff (Mobile)": (12.8, 4.5),
    }
    actor_links = {
        "Admin": ["View Dashboard", "Import Workforce Data", "Configure System", "Audit Activity",
                  "Track Compliance", "Manage Inventory"],
        "Manager": ["View Dashboard", "Manage Schedules", "Monitor Wellness", "Generate Reports"],
        "Analyst": ["View Dashboard", "Analyze Workload", "Train AI Model", "Generate Reports"],
        "Scheduler": ["Manage Schedules", "Mobile Schedule View"],
        "Healthcare\nStaff (Mobile)": ["Mobile Schedule View", "Monitor Wellness"],
    }

    for name, (ax_x, ax_y) in actors.items():
        draw_actor(ax, ax_x, ax_y, name)
        for uc_name in actor_links.get(name, []):
            uc = uc_boxes[uc_name]
            draw_arrow(ax, ax_x + 0.35 if ax_x < 7 else ax_x - 0.35, ax_y,
                       uc["x"] - uc["w"] / 2 - 0.05 if ax_x < 7 else uc["x"] + uc["w"] / 2 + 0.05, uc["y"],
                       color="#78909C", lw=1.0, style="-")

    # Include relationship
    draw_orthogonal(ax, [(5.2, 5.5), (5.75, 5.5), (5.75, 2.9)], color="#EF6C00")
    ax.text(5.5, 5.65, "«include»", fontsize=7, color="#EF6C00", fontstyle="italic")

    return save_fig("usecase_diagram.png")


def draw_sequence():
    fig, ax = plt.subplots(figsize=(15, 8))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Sequence Diagram — AI Workload Model Retraining", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    participants = [
        ("Administrator", 1.2),
        ("Next.js\nFrontend", 3.8),
        ("Spring Boot\nBackend", 6.8),
        ("Python AI\nService", 10.0),
        ("PostgreSQL\nDatabase", 13.0),
    ]

    for name, x in participants:
        ax.plot([x, x], [0.6, 7.0], color="#B0BEC5", lw=1.2, linestyle="--", zorder=0)
        box = FancyBboxPatch((x - 0.85, 7.0), 1.7, 0.65, boxstyle="round,pad=0.04",
                             facecolor=C_LAYER_FE, edgecolor=C_PRIMARY, linewidth=1.5, zorder=2)
        ax.add_patch(box)
        ax.text(x, 7.32, name, ha="center", va="center", fontsize=8, fontweight="medium", zorder=3)

    steps = [
        (1.2, 3.8, 6.5, "1. Click 'Retrain Model'", False),
        (3.8, 6.8, 6.0, "2. POST /api/predictions/retrain  [JWT]", False),
        (6.8, 13.0, 5.4, "3. SELECT workload records", False),
        (13.0, 6.8, 4.8, "4. Return historical data", True),
        (6.8, 10.0, 4.2, "5. POST /train  {data, config}", False),
        (10.0, 6.8, 3.6, "6. Return metrics + model artifact", True),
        (6.8, 13.0, 3.0, "7. INSERT PredictionModel", False),
        (6.8, 10.0, 2.4, "8. POST /forecast", False),
        (10.0, 6.8, 1.8, "9. Return forecast series", True),
        (6.8, 3.8, 1.2, "10. JSON {predictions, metrics}", False),
        (3.8, 1.2, 0.6, "11. Render forecast chart", False),
    ]

    for x1, x2, y, label, dashed in steps:
        draw_arrow(ax, x1, y, x2, y, label=label, dashed=dashed, lw=1.6 if not dashed else 1.2,
                   label_offset=(0, 0.1))
        # Activation bar on backend for steps 2-10
        if x2 == 6.8 or x1 == 6.8:
            bar = FancyBboxPatch((6.55, y - 0.08), 0.5, 0.16, boxstyle="square,pad=0",
                                 facecolor="#A5D6A7", edgecolor="#2E7D32", linewidth=0.8, zorder=1)
            ax.add_patch(bar)

    ax.text(7.5, 0.25, "Solid arrows = request/call   |   Dashed arrows = response/return",
            ha="center", fontsize=8, color=C_LABEL, style="italic")
    return save_fig("sequence_diagram.png")


def draw_system_architecture():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("System Architecture Diagram — Deployment View", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    def layer_box(y, h, label, color):
        rect = FancyBboxPatch((0.5, y), 13.0, h, boxstyle="round,pad=0.06",
                              facecolor=color, edgecolor="#90A4AE", linewidth=1.5, alpha=0.25, zorder=0)
        ax.add_patch(rect)
        ax.text(0.8, y + h - 0.25, label, fontsize=10, fontweight="bold", color=C_SECONDARY)

    layer_box(6.5, 2.0, "Client Tier", C_LAYER_FE)
    layer_box(3.8, 2.3, "Application Tier", C_LAYER_BE)
    layer_box(0.8, 2.5, "Data Tier", C_LAYER_DB)

    nodes = [
        (3.0, 7.3, "Web Browser\nNext.js 16\nPort 3000", C_LAYER_FE, C_PRIMARY),
        (7.0, 7.3, "Mobile App\nExpo / React Native", C_LAYER_FE, C_PRIMARY),
        (11.0, 7.3, "Admin Workstation\nBrowser-based UI", C_LAYER_FE, C_PRIMARY),
        (4.5, 4.7, "Spring Boot 3\nREST API + JWT\nPort 8081", C_LAYER_BE, "#2E7D32"),
        (9.5, 4.7, "Python FastAPI\nML Engines\nPort 8000", C_LAYER_AI, "#EF6C00"),
        (4.0, 1.8, "PostgreSQL 16\nPrimary Data Store\nPort 5432", C_LAYER_DB, "#6A1B9A"),
        (8.5, 1.8, "MongoDB (Optional)\nExtended Audit Logs\nPort 27017", C_LAYER_DB, "#6A1B9A"),
        (12.0, 1.8, "Redis (Optional)\nSession Cache\nPort 6379", C_LAYER_DB, "#6A1B9A"),
    ]

    for x, y, label, fc, ec in nodes:
        box = FancyBboxPatch((x - 1.25, y - 0.55), 2.5, 1.1, boxstyle="round,pad=0.05",
                             facecolor=fc, edgecolor=ec, linewidth=2, zorder=2)
        ax.add_patch(box)
        ax.text(x, y, label, ha="center", va="center", fontsize=7.5, fontweight="medium", zorder=3)

    connections = [
        (3.0, 6.75, 4.5, 5.25, "HTTPS\n/api/* proxy"),
        (7.0, 6.75, 4.5, 5.25, "REST\nJSON + JWT"),
        (11.0, 6.75, 4.5, 5.25, "HTTPS"),
        (4.5, 4.15, 4.0, 2.35, "JDBC\nHibernate/JPA"),
        (4.5, 4.15, 9.5, 5.25, "HTTP\n/train, /forecast"),
        (9.5, 4.15, 4.0, 2.35, "Read\nworkload history"),
        (4.5, 4.15, 8.5, 2.35, "Optional\naudit export"),
    ]
    for x1, y1, x2, y2, label in connections:
        rad = 0.15 if abs(x1 - x2) > 3 else 0
        draw_arrow(ax, x1, y1, x2, y2, label=label, rad=rad, lw=1.8)

    return save_fig("system_architecture.png")


def draw_app_architecture():
    fig, ax = plt.subplots(figsize=(15, 9))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Application Architecture Diagram — Logical Component View", fontsize=15,
                 fontweight="bold", pad=16, color=C_SECONDARY)

    columns = [
        ("Frontend\nNext.js 16", 2.5, C_LAYER_FE, C_PRIMARY, [
            "App Router (16 pages)",
            "React Components + Hooks",
            "API Proxy Middleware",
            "Auth Context (JWT)",
            "Recharts Visualizations",
            "Tailwind CSS UI",
        ]),
        ("Backend\nSpring Boot 3", 7.5, C_LAYER_BE, "#2E7D32", [
            "REST Controllers (20+)",
            "Business Services",
            "JPA Repositories",
            "Security Filter Chain",
            "AiServiceClient (HTTP)",
            "Exception Handler",
        ]),
        ("AI Service\nFastAPI", 12.5, C_LAYER_AI, "#EF6C00", [
            "ml_engine (forecasting)",
            "wellness_engine (burnout)",
            "inventory_engine (demand)",
            "skills_engine (gaps)",
            "scikit-learn models",
            "Pydantic validation",
        ]),
    ]

    for title, cx, fc, ec, items in columns:
        frame = FancyBboxPatch((cx - 2.0, 1.2), 4.0, 6.8, boxstyle="round,pad=0.08",
                               facecolor=fc, edgecolor=ec, linewidth=2, alpha=0.3, zorder=0)
        ax.add_patch(frame)
        ax.text(cx, 7.65, title, ha="center", fontsize=11, fontweight="bold", color=ec)
        for i, item in enumerate(items):
            box = FancyBboxPatch((cx - 1.75, 6.5 - i * 0.72), 3.5, 0.55, boxstyle="round,pad=0.03",
                                 facecolor="white", edgecolor=ec, linewidth=1.2, zorder=2)
            ax.add_patch(box)
            ax.text(cx, 6.77 - i * 0.72, item, ha="center", va="center", fontsize=7.5, zorder=3)

    # Inter-column arrows
    draw_arrow(ax, 4.5, 4.5, 5.5, 4.5, label="REST + JWT\n(JSON over HTTP)", lw=2.2, color=C_ACCENT)
    draw_arrow(ax, 5.5, 4.2, 4.5, 4.2, dashed=True, label="Response DTO", lw=1.3, color=C_LABEL)

    draw_arrow(ax, 9.5, 4.5, 10.5, 4.5, label="HTTP ML API\n(/train, /forecast)", lw=2.2, color=C_ACCENT)
    draw_arrow(ax, 10.5, 4.2, 9.5, 4.2, dashed=True, label="Predictions", lw=1.3, color=C_LABEL)

    # Database
    db = FancyBboxPatch((5.5, 0.3), 4.0, 0.85, boxstyle="round,pad=0.06",
                        facecolor=C_LAYER_DB, edgecolor="#6A1B9A", linewidth=2.5, zorder=2)
    ax.add_patch(db)
    ax.text(7.5, 0.72, "PostgreSQL — Shared Persistence Layer", ha="center", va="center",
            fontsize=9.5, fontweight="bold", color="#6A1B9A")

    draw_arrow(ax, 7.5, 1.15, 7.5, 1.9, label="JPA / JDBC", lw=2.0, color="#6A1B9A")
    draw_orthogonal(ax, [(12.5, 1.2), (12.5, 0.72), (9.5, 0.72)], label="Read training data", color="#6A1B9A")

    return save_fig("app_architecture.png")


# ── Document content ──────────────────────────────────────────────────────────

DIAGRAM_EXPLANATIONS = {
    "erd": {
        "intro": (
            "The Entity Relationship Diagram (ERD) models the PostgreSQL database schema used by HWO. "
            "All tables are managed through JPA/Hibernate entities in the Spring Boot backend and are "
            "auto-created or migrated on application startup."
        ),
        "sections": [
            ("Core Identity Entities", [
                "Department — Represents hospital units (Emergency, ICU, Pediatrics, etc.) with workload aggregates.",
                "Staff — Clinical and administrative workforce records linked to a department.",
                "User — Application login accounts with role-based permissions; optionally linked to a Staff record.",
                "UserProfile — Stores per-user preferences such as dashboard widget visibility and notification settings.",
            ]),
            ("Operational Entities", [
                "Schedule — Daily shift assignments connecting staff to departments with shift type and status.",
                "WorkloadRecord — Hourly patient volume and workload metrics used for analytics and AI training.",
                "WellnessRecord — Staff wellness scores and overtime tracking for burnout detection.",
                "Certification — Professional credentials with expiry dates for compliance monitoring.",
            ]),
            ("Support Entities", [
                "Resource — Inventory items with stock levels, SKU, and reorder thresholds.",
                "ProcurementRequest — Purchase workflow from request through approval to receipt.",
                "AuditLog — Immutable activity trail capturing user actions across all modules.",
                "PredictionModel — Stored ML model metadata, accuracy metrics, and artifact references.",
            ]),
            ("Key Relationships", [
                "Department 1:N Staff — Each staff member belongs to one department; departments aggregate many staff.",
                "Staff 1:N Schedule — A staff member can have multiple shift records over time.",
                "Department 1:N WorkloadRecord — Workload is tracked per department at hourly granularity.",
                "User 1:1 UserProfile — Each authenticated user has exactly one profile record.",
                "User 1:N AuditLog — All significant actions are attributed to the acting user.",
            ]),
        ],
        "caption": "Figure 1: Entity Relationship Diagram — PostgreSQL database schema with primary/foreign key relationships",
    },
    "class": {
        "intro": (
            "The class diagram illustrates the layered architecture of the Spring Boot backend following "
            "separation-of-concerns principles. Each layer has a distinct responsibility and communicates "
            "only with adjacent layers."
        ),
        "sections": [
            ("Controller Layer (Presentation)", [
                "Receives HTTP requests from the Next.js frontend via REST endpoints.",
                "Validates request parameters and delegates to appropriate service classes.",
                "Returns JSON responses or HTTP status codes; protected by JWT security filters.",
                "Examples: PredictionsController handles /api/predictions, ScheduleController handles /api/schedules.",
            ]),
            ("Service Layer (Business Logic)", [
                "Encapsulates all business rules, orchestration, and cross-cutting concerns.",
                "AiServiceClient communicates with the Python FastAPI service for ML operations.",
                "Services are transactional and coordinate between multiple repositories when needed.",
                "Examples: PredictionService trains models; SchedulingService manages shift assignments.",
            ]),
            ("Repository Layer (Persistence)", [
                "Spring Data JPA interfaces providing CRUD and custom query methods.",
                "Abstracts SQL/PostgreSQL access; generates queries from method names or @Query annotations.",
                "One repository per entity type (e.g., StaffRepository, WorkloadRecordRepository).",
            ]),
            ("Entity Layer (Domain Model)", [
                "JPA @Entity classes mapped directly to PostgreSQL tables.",
                "Define relationships (@ManyToOne, @OneToMany) matching the ERD.",
                "Contain getters/setters; no business logic — kept in service layer.",
            ]),
        ],
        "caption": "Figure 2: Class Diagram — Spring Boot layered architecture with dependency flow",
    },
    "usecase": {
        "intro": (
            "The use case diagram defines the functional scope of HWO from the perspective of human actors. "
            "Each actor represents a user role with specific permissions configured in the role matrix."
        ),
        "sections": [
            ("Actors", [
                "Admin — Full system access: configuration, user management, audit, and all operational modules.",
                "Manager — Department-level operations: scheduling, wellness, reporting, and inventory.",
                "Analyst — Data-focused role: workload analysis, AI model training, and report generation.",
                "Scheduler — Shift coordination: schedule management and mobile schedule viewing.",
                "Healthcare Staff (Mobile) — Frontline workers using the mobile app for schedules and wellness.",
            ]),
            ("Primary Use Cases", [
                "View Dashboard — Executive KPI overview available to all roles with menu access.",
                "Import Workforce Data — CSV bulk import of staff, shifts, and workload records.",
                "Analyze Workload — Deep analytics on overtime, ratios, anomalies, and skill mix.",
                "Train AI Model — Trigger ML retraining pipeline for workload forecasting.",
                "Manage Schedules — Create, swap, publish shifts with AI-assisted assignment.",
                "Track Compliance — Run regulatory scans and manage mandate submissions.",
            ]),
            ("Role Restrictions", [
                "Viewer role can only access Dashboard and Reporting (read-only).",
                "Scheduler cannot access AI Prediction or Configuration modules.",
                "Analyst cannot manage users or modify system configuration.",
                "All use cases require JWT authentication; unauthorized access returns HTTP 403.",
            ]),
        ],
        "caption": "Figure 3: Use Case Diagram — Actor interactions with HWO functional modules",
    },
    "sequence": {
        "intro": (
            "This sequence diagram traces the complete AI model retraining workflow — one of the most "
            "complex cross-service interactions in HWO. It demonstrates how four system components "
            "collaborate to train a new workload prediction model and return forecast results."
        ),
        "sections": [
            ("Phase 1: Initiation (Steps 1–2)", [
                "The administrator clicks 'Retrain Model' on the AI Prediction page.",
                "The Next.js frontend sends POST /api/predictions/retrain with the JWT bearer token.",
                "The Spring Boot Security filter validates the token before reaching PredictionsController.",
            ]),
            ("Phase 2: Data Preparation (Steps 3–4)", [
                "PredictionService queries WorkloadRecordRepository for historical workload data.",
                "PostgreSQL returns all department hourly records used as training features.",
                "Data is transformed into the format expected by the Python /train endpoint.",
            ]),
            ("Phase 3: Model Training (Steps 5–7)", [
                "Backend sends workload time-series to Python AI service POST /train.",
                "The ml_engine trains a Ridge + HistGradientBoosting ensemble with CV-tuned weights.",
                "Python returns model metrics (MAPE, RMSE) and a serialized model artifact.",
                "Backend persists a PredictionModel record in PostgreSQL with accuracy and timestamp.",
            ]),
            ("Phase 4: Forecast Generation (Steps 8–11)", [
                "Backend immediately calls POST /forecast with the newly trained model.",
                "Python returns predicted workload values for upcoming periods.",
                "Backend packages predictions as JSON and returns to the frontend.",
                "The frontend renders the forecast chart using Recharts with confidence bands.",
            ]),
        ],
        "caption": "Figure 4: Sequence Diagram — AI workload model retraining end-to-end flow",
    },
    "system": {
        "intro": (
            "The system architecture diagram shows the physical deployment topology of HWO across three tiers. "
            "This view is essential for understanding network boundaries, port assignments, and service dependencies."
        ),
        "sections": [
            ("Client Tier (Port 3000)", [
                "Next.js 16 web application served to desktop browsers.",
                "Expo mobile app for iOS/Android connecting via the same API proxy.",
                "All client requests route through /api/* which Next.js middleware proxies to the backend.",
            ]),
            ("Application Tier (Ports 8081, 8000)", [
                "Spring Boot 3 — Primary API server handling auth, CRUD, business logic, and orchestration.",
                "Python FastAPI — Dedicated ML microservice for predictions, wellness, inventory, and skills AI.",
                "Backend calls AI service via HTTP only when ML features are needed; graceful fallback when offline.",
            ]),
            ("Data Tier (Port 5432+)", [
                "PostgreSQL 16 — Primary relational database for all operational data.",
                "MongoDB (optional) — Extended audit log storage for high-volume environments.",
                "Redis (optional) — Session caching and rate limiting for production deployments.",
            ]),
            ("Communication Protocols", [
                "Client → Backend: HTTPS with JWT bearer tokens in Authorization header.",
                "Backend → PostgreSQL: JDBC connection pool via HikariCP with Hibernate ORM.",
                "Backend → AI Service: HTTP REST calls to /train, /forecast, /analyze endpoints.",
                "Docker Compose orchestrates PostgreSQL container; other services run natively or in containers.",
            ]),
        ],
        "caption": "Figure 5: System Architecture — Three-tier deployment with service ports and protocols",
    },
    "app": {
        "intro": (
            "The application architecture diagram reveals the internal logical structure within each tier. "
            "Unlike the deployment view, this diagram shows software components and their dependencies."
        ),
        "sections": [
            ("Frontend Components (Next.js)", [
                "App Router — File-based routing for 16 module pages under src/app/(app)/.",
                "React Components — Reusable UI elements: charts, tables, forms, modals, pagination.",
                "API Proxy — Next.js rewrites /api/* to Spring Boot backend URL from environment config.",
                "Auth Context — React context provider managing JWT token in localStorage and cookie.",
                "Recharts — Data visualization library for workload trends, forecasts, and KPI charts.",
            ]),
            ("Backend Components (Spring Boot)", [
                "20+ REST Controllers exposing /api/* endpoints with role-based authorization.",
                "Service classes implementing business logic with @Transactional boundaries.",
                "JPA Repositories providing type-safe database access with Spring Data.",
                "Security Filter Chain — JWT validation, CORS, and permission checks on every request.",
                "AiServiceClient — RestTemplate/WebClient wrapper for Python AI service communication.",
            ]),
            ("AI Service Components (FastAPI)", [
                "ml_engine — Workload forecasting with ensemble models and feature engineering.",
                "wellness_engine — Burnout risk scoring and sentiment analysis on staff feedback.",
                "inventory_engine — Demand prediction and reorder optimization for supplies.",
                "skills_engine — Competency gap analysis and training prioritization.",
            ]),
            ("Data Flow Summary", [
                "User action → React page → fetch(/api/...) → Next.js proxy → Spring Controller.",
                "Controller → Service → Repository → PostgreSQL → return entity/DTO up the stack.",
                "For AI features: Service → AiServiceClient → FastAPI → scikit-learn → JSON response.",
            ]),
        ],
        "caption": "Figure 6: Application Architecture — Internal component structure and data flow",
    },
}


SCREENSHOT_EXPLANATIONS = [
    {
        "file": "01-analytics-dashboard.png",
        "title": "Analytics Dashboard",
        "route": "/dashboard",
        "purpose": "Data-driven workforce planning hub with KPIs, trends, forecasts, and staffing shortage alerts.",
        "description": (
            "The Analytics Dashboard aggregates workload, wellness, AI predictions, and scheduling coverage "
            "into a single decision-support view. It supports workforce planners with overtime trends, "
            "burnout risk distribution, utilization rates, and future staffing forecasts."
        ),
        "elements": [
            "KPI Cards — Total staff, average workload, AI prediction accuracy (MAPE/R²), wellness alert count.",
            "Department Workload Trends — Bar chart comparing workload intensity across units.",
            "Peak Workload by Hour — 24-hour demand profile for capacity planning.",
            "Actual vs Predicted Trend — ML forecast overlay for proactive staffing.",
            "Overtime Trend — Monthly average overtime from wellness records.",
            "Burnout Risk Distribution — Count of low/medium/high risk staff.",
            "Schedule Coverage & Shortages — Coverage % and AI staffing recommendations.",
            "Future Staffing Forecast — AI-predicted workload for upcoming periods.",
            "Wellness Alerts — At-risk staff with overtime and risk level.",
            "Staff Allocation Heatmap — Department × shift intensity grid.",
        ],
        "apis": [
            "GET /api/dashboard/analytics", "GET /api/departments", "GET /api/predictions",
            "GET /api/wellness", "GET /api/workload?type=byHour|trend", "GET /api/dashboard/heatmap",
        ],
        "users": "Hospital administrators, workforce planners, and department heads.",
    },
    {
        "file": "02-data-collection.png",
        "title": "Data Collection",
        "route": "/data-collection",
        "purpose": "Central hub for importing workforce data and managing staff records.",
        "description": (
            "This module is the data ingestion gateway for HWO. All downstream analytics, AI predictions, "
            "and scheduling depend on clean data imported here via CSV templates or manual entry."
        ),
        "elements": [
            "Summary KPIs — Staff records count, schedule rows, workload records, and last import quality score.",
            "Import Tab — Three CSV templates: Staff Roster, Shift Schedule, and Patient/Workload Volume.",
            "Template Cards — Expandable cards showing required columns with Download CSV and bulk sample options.",
            "Upload Area — Drag-and-drop CSV upload with type selector (staff/shift/patient).",
            "Validation Summary — Valid count, duplicates, errors, and quality percentage from last import.",
            "Staff Tab — Manual staff entry and searchable roster management.",
            "History Tab — Paginated import run history with date filtering.",
            "Scheduler Tab — HIS/HR sync frequency configuration.",
        ],
        "apis": ["POST /api/import", "GET /api/import/templates/{type}", "GET /api/staff", "POST /api/staff", "GET /api/import/history"],
        "users": "Data administrators and HR staff responsible for workforce records.",
    },
    {
        "file": "03-workload-analysis.png",
        "title": "Workload Analysis",
        "route": "/workload-analysis",
        "purpose": "Deep-dive analytics for staffing load, overtime, ratios, and anomalies.",
        "description": (
            "While the Dashboard provides a high-level overview, Workload Analysis offers granular insights "
            "for workforce analysts to identify bottlenecks, compliance risks, and optimization opportunities."
        ),
        "elements": [
            "Summary KPIs — Peak load hour, average nurse-to-patient ratio, total overtime hours, anomaly count.",
            "Hourly Workload Chart — Demand distribution across 24 hours for capacity planning.",
            "Monthly Trend Chart — Historical workload patterns for seasonal planning.",
            "Overtime by Department — Bar chart highlighting units with excessive overtime.",
            "Nurse-to-Patient Ratios — Compliance-style ratio tracking per department.",
            "Skill Mix Analysis — Role distribution visualization per department.",
            "Anomaly Detection — Flagged unusual workload spikes requiring investigation.",
            "Staff Wellness Overlay — Filterable staff list with wellness indicators by department.",
        ],
        "apis": ["GET /api/workload/summary", "GET /api/workload/overtime", "GET /api/workload/ratios", "GET /api/workload/anomalies", "GET /api/workload/skill-mix"],
        "users": "Workforce analysts, nursing directors, and operations managers.",
    },
    {
        "file": "04-ai-prediction.png",
        "title": "AI Prediction & Forecasting",
        "route": "/ai-prediction",
        "purpose": "Workforce demand forecasting, staff requirement prediction, and explainable ML outputs.",
        "description": (
            "Core AI module using Ridge + HistGradientBoosting ensemble (ml_engine) for workload forecasting. "
            "Displays model evaluation metrics (MAE, RMSE, R², CV-MAE), feature importance, department-level "
            "daily models for scheduling surge targets, and recommended staffing levels derived from forecasts."
        ),
        "elements": [
            "Retrain Model — Trains global monthly + per-department daily models via POST /api/predictions/retrain.",
            "Forecast Chart — Predicted workload with confidence bands (low/high).",
            "Model Metadata — Algorithm type, accuracy (100−MAPE), last trained date, training data points.",
            "Evaluation Metrics — MAE, RMSE, R², baseline comparison, improvement vs naive forecast.",
            "Feature Importance — Top drivers (lag features, rolling stats, seasonality) with importance scores.",
            "Model Comparison — Side-by-side metrics for two saved model versions.",
            "Staffing Recommendations — Links forecast multipliers to scheduling effectiveMinStaff targets.",
            "Export — Download predictions CSV for external planning tools.",
        ],
        "apis": [
            "POST /api/predictions/retrain", "GET /api/predictions", "GET /api/predictions/models",
            "GET /api/predictions/compare", "GET /api/scheduling/ai/forecast",
        ],
        "users": "Data scientists, workforce planners, and administrators.",
    },
    {
        "file": "05-scheduling.png",
        "title": "Scheduling Optimization",
        "route": "/scheduling",
        "purpose": "AI-assisted schedule optimization with conflict detection, fair distribution, and staffing recommendations.",
        "description": (
            "Demonstrates how AI contributes to workforce optimization — not merely a calendar tool. "
            "Uses forecast-driven surge targets, ranked assignee suggestions (40% rules + 60% AI blend), "
            "automatic conflict detection, swap workflow, and what-if scenario planning."
        ),
        "elements": [
            "Coverage KPIs — Schedule coverage %, open shifts, pending swap requests.",
            "AI Auto-Schedule — Fills gaps using AI-ranked staff; reports coverage before/after.",
            "Conflict Detection — Double-booking, leave overlap, preference violations, skill-mix gaps.",
            "Forecast-Driven Targets — Department cards showing predicted load and effectiveMinStaff.",
            "Assignee Suggestions — Scored staff list with reasons (wellness risk, certs, AI rank).",
            "Swap Approval Workflow — Request swap from schedule table rows.",
            "Constraints Editor — Max hours, rest between shifts, fairness rules.",
            "What-If Scenarios — Project coverage impact of adding shifts.",
            "Overtime Monitoring — Flags staff approaching hour limits.",
        ],
        "apis": [
            "POST /api/scheduling/ai/auto-schedule", "GET /api/scheduling/conflicts",
            "GET /api/scheduling/ai/suggestions", "GET /api/scheduling/ai/forecast",
            "POST /api/scheduling/ai/what-if", "POST /api/schedules/swap",
        ],
        "users": "Scheduling coordinators, charge nurses, and department managers.",
    },
    {
        "file": "07-reporting.png",
        "title": "Reporting",
        "route": "/reporting",
        "purpose": "Generate, customize, and schedule workforce reports for stakeholders.",
        "description": (
            "The Reporting module transforms operational data into formatted reports suitable for "
            "executive briefings, compliance submissions, and departmental reviews."
        ),
        "elements": [
            "Report Library — Pre-built report types (workload summary, wellness, scheduling, compliance).",
            "Generate Button — One-click report generation for standard types.",
            "Benchmark Comparison — Hospital metrics vs. industry benchmark values.",
            "Custom Report Builder — Select metrics and date range for ad-hoc reports.",
            "Executive Summary — High-level narrative report from current workforce data.",
            "Scheduled Reports — Configure recurring report delivery (daily, weekly, monthly).",
            "Download/View — Access generated report outputs.",
        ],
        "apis": ["GET /api/reports", "POST /api/reports/generate", "GET /api/reports/benchmarks", "POST /api/reports/schedule"],
        "users": "Executives, compliance officers, and analysts.",
    },
    {
        "file": "06-wellness.png",
        "title": "Staff Wellness & Burnout Prediction",
        "route": "/wellness",
        "purpose": "Explainable burnout risk prediction with feature contributions and actionable interventions.",
        "description": (
            "Addresses the black-box concern with transparent AI outputs. Uses HistGradientBoostingClassifier "
            "with published metrics (accuracy, precision, recall, F1, ROC-AUC). Each alert shows WHY an "
            "employee was flagged with contributing factor percentages and recommended actions."
        ),
        "elements": [
            "AI Methodology Panel — Model name, algorithm, evaluation metrics (F1, ROC-AUC, precision, recall).",
            "Burnout Risk Alerts — Risk level + probability with 'Why flagged?' explanation text.",
            "Contributing Factors — e.g. Excessive overtime (+35%), Consecutive night shifts (+25%), Low wellness score (+20%).",
            "Recommended Actions — Per-factor and top-pick interventions (reduce overtime, schedule recovery, manager review).",
            "Wellness Trend Chart — Score trajectory over time.",
            "Daily Check-in — Regular wellness submissions feeding the model.",
            "Staff Satisfaction Survey — Structured feedback collection.",
            "Intervention Assignment — One-click assign AI-recommended or manual interventions.",
        ],
        "apis": [
            "GET /api/wellness", "GET /api/wellness/ai/model-info", "GET /api/wellness/ai/risk/{staffId}",
            "POST /api/wellness/checkin", "POST /api/wellness/interventions",
        ],
        "users": "HR teams, employee health coordinators, and managers.",
    },
    {
        "file": "08-mobile-schedule.png",
        "title": "Mobile — Schedule View",
        "platform": "Expo (React Native) · iOS/Android",
        "meta_label": "Platform",
        "mobile": True,
        "purpose": "Native mobile access for frontline staff to view shifts and request swaps.",
        "description": (
            "The HWO mobile app is built with Expo SDK 52 and React Native (not a web configuration page). "
            "It connects to the same Spring Boot API via EXPO_PUBLIC_API_URL. Supports offline schedule "
            "caching via AsyncStorage when connectivity is limited."
        ),
        "elements": [
            "7-Day Schedule List — Shifts pulled from GET /api/mobile/schedules with department and status.",
            "Shift Swap Request — In-app swap workflow calling POST /api/schedules/swap.",
            "Offline Mode — Cached schedule displayed when network unavailable.",
            "Tab Navigation — Schedule, Wellness, Alerts (Expo Router file-based routing).",
            "Secure Auth — JWT via POST /api/auth/mobile-login stored in SecureStore.",
        ],
        "apis": ["GET /api/mobile/schedules", "POST /api/schedules/swap", "POST /api/auth/mobile-login"],
        "users": "Frontline healthcare staff (nurses, clinicians).",
    },
    {
        "file": "09-mobile-wellness.png",
        "title": "Mobile — Wellness Check-in",
        "platform": "Expo (React Native) · iOS/Android",
        "meta_label": "Platform",
        "mobile": True,
        "purpose": "Daily wellness check-ins and burnout risk visibility for staff on the go.",
        "description": (
            "Mobile wellness tab shows personal wellness score, risk badge, and overtime hours from "
            "GET /api/mobile/wellness. Staff submit daily check-ins that feed the burnout prediction pipeline."
        ),
        "elements": [
            "Wellness Score Ring — Personal score out of 100.",
            "Risk Badge — Low/medium/high from AI burnout classifier.",
            "Overtime Display — Weekly overtime hours contributing to risk.",
            "Daily Check-in Button — Submits wellness data to backend.",
        ],
        "apis": ["GET /api/mobile/wellness", "POST /api/wellness/checkin"],
        "users": "Frontline healthcare staff.",
    },
    {
        "file": "10-mobile-alerts.png",
        "title": "Mobile — Real-time Alerts",
        "platform": "Expo (React Native) · iOS/Android",
        "meta_label": "Platform",
        "mobile": True,
        "purpose": "Push-style notifications for schedule changes, wellness reminders, and workload surges.",
        "description": (
            "Alerts tab displays workforce notifications from GET /api/mobile/alerts including schedule "
            "changes, wellness reminders, and high-workload department warnings linked to AI forecasts."
        ),
        "elements": [
            "Schedule Change Alerts — Updated shift times or departments.",
            "Wellness Reminders — Prompt to complete daily check-in.",
            "Workload Surge Alerts — Forecast-driven staffing notifications.",
            "Timestamped Feed — Chronological alert history.",
        ],
        "apis": ["GET /api/mobile/alerts"],
        "users": "Frontline healthcare staff.",
    },
]


# ── Document builders ─────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_paragraphs(doc, text):
    for para in text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())


def add_image(doc, path: Path, caption: str, width=Inches(6.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Image not found: {path.name}]")


def add_diagram_section(doc, key: str, diagram_path: Path):
    info = DIAGRAM_EXPLANATIONS[key]
    add_paragraphs(doc, info["intro"])
    doc.add_paragraph()
    add_image(doc, diagram_path, info["caption"])
    doc.add_paragraph()
    add_heading(doc, "Diagram Explanation", level=3)
    for section_title, bullets in info["sections"]:
        add_heading(doc, section_title, level=4)
        add_bullets(doc, bullets)
        doc.add_paragraph()


def add_screenshot_section(doc, info: dict):
    add_heading(doc, info["title"], level=2)
    p = doc.add_paragraph()
    meta = info.get("route") or info.get("platform", "")
    run = p.add_run(f"{info.get('meta_label', 'Route')}: {meta}  |  ")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x71, 0x64)
    run2 = p.add_run(f"Primary Users: {info['users']}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x71, 0x64)

    add_heading(doc, "Purpose", level=3)
    doc.add_paragraph(info["purpose"])
    add_heading(doc, "Overview", level=3)
    doc.add_paragraph(info["description"])

    add_heading(doc, "Key UI Elements", level=3)
    add_bullets(doc, info["elements"])

    add_heading(doc, "Backend APIs", level=3)
    add_bullets(doc, info["apis"])

    add_heading(doc, "Screenshot", level=3)
    caption = f"Figure: {info['title']}" + (f" — {info.get('route', info.get('platform', ''))}" if info.get('route') or info.get('platform') else "")
    width = Inches(3.2) if info.get("mobile") else Inches(6.3)
    add_image(doc, SCREENSHOT_DIR / info["file"], caption, width=width)
    doc.add_paragraph()


def build_document():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)

    print("Generating improved diagrams...")
    diagrams = {
        "erd": draw_erd(),
        "class": draw_class_diagram(),
        "usecase": draw_usecase(),
        "sequence": draw_sequence(),
        "system": draw_system_architecture(),
        "app": draw_app_architecture(),
    }

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Health Workforce Optimizer (HWO)\n")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    sub = title.add_run("\nProject Documentation\n")
    sub.font.size = Pt(18)
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Sections 1-5 (unchanged content, already detailed)
    add_heading(doc, "1. General Description")
    doc.add_paragraph(
        "The Health Workforce Optimizer (HWO) is an AI-based healthcare workforce management system "
        "focused on core workforce optimization: data collection, workload analytics, AI-driven forecasting, "
        "intelligent scheduling, explainable burnout prediction, analytics dashboards, reporting, and mobile "
        "workforce access. Non-core modules (inventory management, general compliance management) are "
        "out of scope for this project documentation."
    )
    doc.add_paragraph(
        "HWO serves hospital administrators, workforce planners, schedulers, analysts, and frontline staff "
        "through a Next.js web application and a native Expo (React Native) mobile app."
    )
    add_heading(doc, "Core System Modules", level=2)
    add_bullets(doc, [
        "Data Collection — CSV import and staff roster management",
        "Workload Analysis — Deep-dive staffing analytics and anomaly detection",
        "AI Prediction & Forecasting — ML workload demand and staffing level forecasts",
        "Scheduling Optimization — AI-assisted assignment, conflict detection, fair distribution",
        "Staff Wellness & Burnout Prediction — Explainable AI risk scoring with interventions",
        "Analytics Dashboard — Workforce KPIs, trends, forecasts, and shortage alerts",
        "Reporting — Executive and operational workforce reports",
        "Mobile Workforce Access — Expo React Native app for schedules, wellness, alerts",
    ])

    add_heading(doc, "2. Problem Statement")
    doc.add_paragraph(
        "Healthcare organizations face persistent challenges in balancing patient demand with available "
        "staffing resources. Manual scheduling, fragmented data sources, and reactive decision-making "
        "lead to several critical problems:"
    )
    add_bullets(doc, [
        "Unpredictable patient volumes cause understaffing during peak hours and overstaffing during quiet periods",
        "High overtime and burnout among clinical staff increase turnover and reduce care quality",
        "Workforce data is scattered across HR systems, shift rosters, and patient census records",
        "Black-box wellness scoring erodes trust — staff see risk alerts without understanding why",
        "Lack of data-driven forecasting prevents proactive workforce planning",
        "Manual scheduling cannot optimize fair workload distribution or detect conflicts early",
    ])
    doc.add_paragraph(
        "These problems result in increased operational costs, staff dissatisfaction, compliance risks, "
        "and suboptimal patient care delivery. HWO addresses these gaps by providing an integrated, "
        "AI-powered platform for workforce optimization."
    )

    add_heading(doc, "3. Proposed Solution")
    doc.add_paragraph("HWO delivers an AI-enhanced workforce optimization platform organized around eight core modules:")
    add_bullets(doc, [
        "Data Collection: CSV import for staff, shifts, and workload records",
        "Workload Analysis: Hourly trends, overtime, ratios, anomalies, skill mix",
        "AI Prediction & Forecasting: Ridge + HistGradientBoosting ensemble for demand and staffing forecasts",
        "Scheduling Optimization: AI auto-schedule, conflict detection, forecast-driven surge targets",
        "Staff Wellness & Burnout Prediction: Explainable GBM classifier with factor breakdowns",
        "Analytics Dashboard: Overtime trends, burnout distribution, utilization, staffing shortages",
        "Reporting: Executive summaries, benchmarks, scheduled delivery",
        "Mobile Workforce Access: Expo React Native app (Schedule, Wellness, Alerts tabs)",
    ])

    add_heading(doc, "3.1 AI/ML Models and Explainability", level=2)
    doc.add_paragraph(
        "HWO uses scikit-learn models served by a Python FastAPI microservice. All predictions include "
        "evaluation metrics and, where applicable, feature importance or contributing-factor breakdowns."
    )
    add_heading(doc, "Workload Forecasting (ml_engine)", level=3)
    add_bullets(doc, [
        "Models: Ridge regression + HistGradientBoostingRegressor ensemble with CV-tuned blend weights",
        "Purpose: Workforce demand forecasting and recommended staffing levels (effectiveMinStaff multipliers)",
        "Metrics: MAE, RMSE, R², CV-MAE, improvement vs naive baseline",
        "Explainability: Feature importance (lag-1/7/14, rolling stats, EWMA, seasonality coefficients)",
        "Output: Monthly global forecast + per-department daily models for scheduling surge targets",
    ])
    add_heading(doc, "Burnout Risk Prediction (wellness_engine)", level=3)
    add_bullets(doc, [
        "Model: HistGradientBoostingClassifier on overtime, wellness score, weekly hours, night shifts, shift irregularity",
        "Metrics: Accuracy, precision, recall, F1-score, ROC-AUC (validation set)",
        "Explainability: Feature-ablation contributions (SHAP-style approximation) + domain rule factors",
        "Example output: 'Burnout Risk: High (82%) — Excessive overtime (+35%), Consecutive night shifts (+25%), Low wellness score (+20%)'",
        "Actions: Per-factor recommended actions (reduce overtime, schedule recovery, manager wellness review)",
    ])
    add_heading(doc, "Scheduling AI (rank-assignees + rules)", level=3)
    add_bullets(doc, [
        "Assignee ranking: Weighted blend of rule score (40%) and AI rank (60%)",
        "Inputs: Department match, preferences, wellness risk, rest compliance, skill match, hours headroom",
        "Auto-schedule: Fills open shifts using ranked assignees against forecast-driven staffing targets",
        "Conflict detection: Double-booking, leave overlap, preference violations, skill-mix gaps",
    ])
    doc.add_paragraph(
        "Architecture: Next.js → Spring Boot → PostgreSQL; Spring Boot → FastAPI for ML inference."
    )

    add_heading(doc, "4. Data Collection Methodology and Tools Used")
    add_heading(doc, "4.1 Data Collection Methods", level=2)
    add_bullets(doc, [
        "CSV File Import: Bulk import of staff, shift, and patient/workload records via standardized templates",
        "Manual Data Entry: Direct staff roster management through the Data Collection module",
        "System Integration: Configurable HIS/HR sync scheduler for automated data refresh",
        "Operational Data Generation: Backend seed data and sample datasets (up to 20,000 rows) for testing",
        "User Activity Logging: Automatic audit trail capture for all significant system actions",
        "Wellness Surveys: Staff feedback and survey responses collected via web and mobile interfaces",
    ])
    add_heading(doc, "4.2 Data Types Collected", level=2)
    add_bullets(doc, [
        "Staff records: name, email, role, department assignment",
        "Shift schedules: staff assignments, dates, shift types, status",
        "Workload/patient volume: hourly department workload, patient counts, staff on duty",
        "Wellness metrics: wellness scores, overtime hours, feedback text, survey responses",
        "Inventory data: stock levels, movements, transfers, procurement requests",
        "Certifications and training: expiry dates, enrollment status, competency records",
        "Compliance records: regulatory submissions, scan results, template configurations",
    ])
    add_heading(doc, "4.3 Tools and Technologies", level=2)
    add_bullets(doc, [
        "Frontend: Next.js 16, React, TypeScript, Tailwind CSS, Recharts",
        "Backend API: Spring Boot 3, Java 17, Spring Security (JWT), Spring Data JPA",
        "AI/ML Service: Python 3.9+, FastAPI, scikit-learn, NumPy, Pandas",
        "Database: PostgreSQL 16",
        "Mobile: Expo (React Native)",
        "DevOps: Docker Compose (PostgreSQL), Maven, npm",
        "Optional: MongoDB (audit logs), Redis (caching)",
    ])

    add_heading(doc, "5. Requirements")
    add_heading(doc, "5.1 Functional Requirements", level=2)
    add_bullets(doc, [
        "FR-01: User authentication and role-based authorization (Admin, Manager, Analyst, Scheduler, Viewer)",
        "FR-02: Dashboard with KPIs, workload charts, wellness alerts, and department allocation heatmap",
        "FR-03: CSV import for staff, shift, and workload data with validation and duplicate detection",
        "FR-04: Workload analysis with hourly trends, overtime, ratios, skill mix, and anomaly detection",
        "FR-05: AI model training and workload forecasting with model comparison and export",
        "FR-06: Schedule management with AI-assisted assignment, swap requests, leave, and on-call",
        "FR-07: Report generation with scheduling and export capabilities",
        "FR-08: Staff wellness monitoring with burnout prediction and intervention recommendations",
        "FR-09: Inventory management with AI demand forecasting and procurement workflow",
        "FR-10: Skills and certification tracking with training program management",
        "FR-11: Compliance scanning, template management, and regulatory submission",
        "FR-12: User management with configurable role permissions",
        "FR-13: System configuration for departments, roles, integrations, and AI settings",
        "FR-14: Data management with purge and maintenance operations",
        "FR-15: Audit logging with search and export",
        "FR-16: User profile management with dashboard preferences",
        "FR-17: Mobile app access for schedules, wellness, and alerts",
    ])
    add_heading(doc, "5.2 Non-Functional Requirements", level=2)
    add_bullets(doc, [
        "NFR-01: System shall support 20,000+ staff records with paginated API responses",
        "NFR-02: API responses shall be secured with JWT authentication",
        "NFR-03: AI service shall gracefully degrade when offline (rule-based fallbacks)",
        "NFR-04: Web UI shall be responsive and accessible across desktop browsers",
        "NFR-05: Database schema shall auto-migrate on backend startup",
        "NFR-06: System shall log all critical operations for audit compliance",
        "NFR-07: CSV imports shall validate data quality and report error statistics",
    ])
    add_heading(doc, "5.3 Hardware/Software Requirements", level=2)
    add_bullets(doc, [
        "Node.js 20+, Java 17+, Python 3.9+, PostgreSQL 16",
        "Minimum 8 GB RAM for running all services concurrently",
        "Modern web browser (Chrome, Firefox, Safari, Edge)",
        "Network access between frontend (:3000), backend (:8081), AI service (:8000), and database (:5432)",
    ])

    # Diagram sections with detailed explanations
    diagram_sections = [
        ("6. Database Diagram (ERD)", "erd"),
        ("7. Class Diagram", "class"),
        ("8. Use Case Diagram", "usecase"),
        ("9. Sequence Diagram", "sequence"),
        ("10. System Architecture Diagram", "system"),
        ("11. Application Architecture Diagram", "app"),
    ]
    for heading, key in diagram_sections:
        add_heading(doc, heading)
        add_diagram_section(doc, key, diagrams[key])
        doc.add_page_break()

    # Screenshot sections with detailed explanations
    add_heading(doc, "12. Frontend UI Screenshots — Core Modules")
    doc.add_paragraph(
        "Screenshots cover the eight core HWO modules plus three native mobile app screens (Expo React Native). "
        "Inventory and compliance modules are excluded from project scope."
    )
    doc.add_paragraph()

    for info in SCREENSHOT_EXPLANATIONS:
        add_screenshot_section(doc, info)

    doc.save(str(OUTPUT))
    print(f"Document saved to: {OUTPUT}")
    missing = [s["file"] for s in SCREENSHOT_EXPLANATIONS if not (SCREENSHOT_DIR / s["file"]).exists()]
    if missing:
        print(f"Missing screenshots ({len(missing)}): {', '.join(missing)}")


if __name__ == "__main__":
    build_document()
